from flask import Flask, request, jsonify
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFDirectoryLoader, PyPDFLoader
from langchain_core.documents import Document
from langchain.embeddings.base import Embeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer, CrossEncoder
from torch.cuda import is_available
from groq import Groq
import os
import logging
import base64
import re
from flask_cors import CORS
from paddleocr import PaddleOCR
import cv2
import numpy as np
import fitz
from sklearn.feature_extraction.text import TfidfVectorizer
import io
from PIL import Image
import json
from docx import Document as DocxDocument
import pandas as pd
from dotenv import load_dotenv
import csv

logging.getLogger('ppocr').setLevel(logging.ERROR)

load_dotenv()
app = Flask(__name__)
CORS(app)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Now configure Groq client
client = Groq(api_key=GROQ_API_KEY)
# Global settings
strict_conversation_history = {}
flexible_conversation_history = {}
# Define a custom directory for PaddleOCR models
model_dir = "D:\\Python_Project\\Ollama_ChatBot2\\paddleocr_models"

# Ensure the directory exists
os.makedirs(model_dir, exist_ok=True)

# Initialize PaddleOCR with the custom directory
ocr = PaddleOCR(use_angle_cls=True,lang="en")
# ocr = PaddleOCR(use_angle_cls=True, lang="en", use_gpu=is_available())

# Prompt templates with dynamic word limit
def get_prompt_template(is_strict=True):
    base_template = """
    Answer questions using only the provided context and conversation history:
    - Provide responses that match the complexity of the question - be concise for simple queries 
      but allow more detail for complex questions when necessary.
    When answering questions about prices or numerical values, only include currency symbols or units if they are explicitly mentioned in the provided context or conversation history; otherwise, provide the number as is.
    - Format responses using bullet points when:
      * Explaining multi-step processes
      * Listing items/requirements/features
      * Comparing multiple options
      * Presenting chronological sequences
    - Use paragraph format for:
      * Simple factual answers
      * Single-concept explanations
      * Direct answers to yes/no questions
    If the question is unrelated to the context, respond with: "I'm sorry, but I don't have information regarding that. Could you ask something related to the provided information?"
    - For greetings or non-questions, give a short, friendly reply.
    - Not include phrases like "friendly response," "main topics," or "context", "Document", "documents", "Document provided" in the answer.
    - After answering, add: "Feel free to ask more!"
    
    <context>
    {context}
    </context>
    <history>
    {history}
    </history>
    Question: {input}
    """
    flexible_template = """
    Answer questions using only the provided context and conversation history:
    - Provide responses that match the complexity of the question - be concise for simple queries 
      but allow more detail for complex questions when necessary.
    When answering questions about prices or numerical values, only include currency symbols or units if they are explicitly mentioned in the provided context or conversation history; otherwise, provide the number as is.
    - Format responses using bullet points when:
      * Explaining multi-step processes
      * Listing items/requirements/features
      * Comparing multiple options
      * Presenting chronological sequences
    - Use paragraph format for:
      * Simple factual answers
      * Single-concept explanations
      * Direct answers to yes/no questions
    If the question is unrelated to the context but can be answered with general knowledge:
    - Present information assertively without hedging
    - Use formal but accessible language
    - Include relevant dates/terms of service
    - Maintain neutral tone
    - Add: "Feel free to ask more!"
    For greetings or non-questions, give a short, friendly reply.
    - Not include phrases like "friendly response," "main topics," or "context", "Document", "documents", "Document provided" in the answer.
    <context>
    {context}
    </context>
    <history>
    {history}
    </history>
    Question: {input}
    """
    
    return ChatPromptTemplate.from_template(base_template if is_strict else flexible_template)

# Custom Embeddings class with tabular data enhancement
class LocalEmbeddings(Embeddings):
    def __init__(self, model_name: str = "all-MiniLM-L12-v2"):
        device = "cuda" if is_available() else "cpu"
        self.model = SentenceTransformer(model_name, device=device)

    def embed_documents(self, texts):
        processed_texts = []
        for text in texts:
            if ", " in text:  # Heuristic for tabular data
                lines = text.split("\n")
                column_names = " ".join([line.split(", ")[0].split(":")[0] for line in lines if line])
                processed_texts.append(column_names + " " + text)
            else:
                processed_texts.append(text)
        return self.model.encode(processed_texts, convert_to_tensor=False, show_progress_bar=True)

    def embed_query(self, text):
        return self.model.encode(text, convert_to_tensor=False, show_progress_bar=True)

# Initialize re-ranking model
cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2')

# Updated dynamic chunking with tabular data support
def dynamic_chunking(docs, rows_per_chunk=10):
    all_chunks = []
    for doc in docs:
        lines = doc.page_content.split("\n")
        if doc.metadata["source"].endswith(('.csv', '.xlsx')) and ", " in lines[0]:
            header = lines[0] if "Sheet:" not in lines[0] else ""
            data_lines = [line for line in lines if line.strip() and "Sheet:" not in line]
            for i in range(0, len(data_lines), rows_per_chunk):
                chunk_lines = data_lines[i:i + rows_per_chunk]
                chunk_text = header + "\n" + "\n".join(chunk_lines) if header else "\n".join(chunk_lines)
                new_doc = Document(
                    page_content=chunk_text,
                    metadata=doc.metadata.copy()
                )
                new_doc.metadata["chunk"] = i // rows_per_chunk + 1
                all_chunks.append(new_doc)
        else:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=100,
                separators=["\n\n", "\n", ". ", " ", ""],
                length_function=len
            )
            chunks = text_splitter.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                new_doc = Document(
                    page_content=chunk,
                    metadata=doc.metadata.copy()
                )
                new_doc.metadata["chunk"] = i + 1
                all_chunks.append(new_doc)
    return all_chunks

def extract_text_from_images(pdf_path):
    doc = fitz.open(pdf_path)
    extracted_text = []
    for page in doc:
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            image_np = np.array(image)
            result = ocr.ocr(image_np, cls=True)
            page_text = []
            if result is not None:
                for line in result:
                    if line:
                        for word_info in line:
                            page_text.append(word_info[1][0])
            extracted_text.append("\n".join(page_text))
    return "\n".join(extracted_text)

@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    data = request.get_json(force=True)
    user_code = data.get("User_Code")
    token = data.get("Token")
    base64_file = data.get("pdf_base64")
    filename = data.get("filename")
    file_extension = data.get("file_extension")
    word_limit = data.get("wordingLength", 50)

    if not all([user_code, token, base64_file, filename, file_extension]):
        return jsonify({"status": "error", "message": "Missing required fields"})

    if not file_extension.startswith('.'):
        file_extension = f".{file_extension}"

    supported_extensions = ['.pdf', '.docx', '.xlsx', '.csv','.txt']
    if file_extension.lower() not in supported_extensions:
        return jsonify({"status": "error", "message": f"Unsupported file type. Supported: {', '.join(supported_extensions)}"})

    user_dir = os.path.join("./clients", user_code)
    token_dir = os.path.join(user_dir, token)
    file_directory = os.path.join(token_dir, "pdfs")
    faiss_indices_dir = os.path.join(token_dir, "faiss_indices")

    os.makedirs(file_directory, exist_ok=True)
    os.makedirs(faiss_indices_dir, exist_ok=True)

    full_filename = f"{filename}{file_extension}"
    file_path = os.path.join(file_directory, full_filename)

    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            faiss_file_dir = os.path.join(faiss_indices_dir, os.path.splitext(full_filename)[0])
            if os.path.exists(faiss_file_dir):
                for f in os.listdir(faiss_file_dir):
                    os.remove(os.path.join(faiss_file_dir, f))
                os.rmdir(faiss_file_dir)

        with open(file_path, "wb") as file:
            file.write(base64.b64decode(base64_file))
        logging.info(f"File saved at: {file_path}")

        split_docs = []
        if file_extension.lower() == '.pdf':
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                image_texts = []
                for img in page.get_images(full=True):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image = Image.open(io.BytesIO(base_image["image"]))
                    result = ocr.ocr(np.array(image), cls=True)
                    if result:
                        image_texts.extend(word_info[1][0] for line in result if line for word_info in line)
                combined_text = page_text + ("\n" + "\n".join(image_texts) if image_texts else "")
                if combined_text.strip():
                    split_docs.append(Document(page_content=combined_text, metadata={"source": filename, "page": page_num + 1}))

        elif file_extension.lower() == '.docx':
            doc = DocxDocument(file_path)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            combined_text = "\n".join(full_text)
            if combined_text.strip():
                split_docs.append(Document(page_content=combined_text, metadata={"source": filename, "page": 1}))

        elif file_extension.lower() == '.xlsx':
            xlsx_data = pd.read_excel(file_path, engine='openpyxl', sheet_name=None)
            full_text = []
            for sheet_name, df in xlsx_data.items():
                df = df.fillna('')
                sheet_text = f"Sheet: {sheet_name}\n"
                for _, row in df.iterrows():
                    row_text = ", ".join([f"{col}: {val}" for col, val in row.items()])
                    sheet_text += row_text + "\n"
                full_text.append(sheet_text)
            combined_text = "\n\n".join(full_text)
            if combined_text.strip():
                split_docs.append(Document(
                    page_content=combined_text,
                    metadata={"source": filename, "page": 1, "columns": ", ".join(df.columns)}
                ))

        elif file_extension.lower() == '.csv':
            # csv_data = pd.read_csv(file_path)

            try:
                csv_data = pd.read_csv(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                csv_data = pd.read_csv(file_path, encoding='latin1')
            csv_data = csv_data.fillna('')
            csv_text = ""
            for _, row in csv_data.iterrows():
                row_text = ", ".join([f"{col}: {val}" for col, val in row.items()])
                csv_text += row_text + "\n"
            if csv_text.strip():
                split_docs.append(Document(
                    page_content=csv_text,
                    metadata={"source": filename, "page": 1, "columns": ", ".join(csv_data.columns)}
                ))
        elif file_extension.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                try:
                    text_content = file.read()
                except UnicodeDecodeError:
                    with open(file_path, 'r', encoding='latin1') as fallback_file:
                        text_content = fallback_file.read()
            if text_content.strip():
                split_docs.append(Document(
                    page_content=text_content,
                    metadata={"source": filename, "page": 1}
                ))


        if not split_docs:
            raise ValueError("No valid content extracted")

        chunked_docs = dynamic_chunking(split_docs)
        if not chunked_docs:
            raise ValueError("No valid chunks created")

        embeddings = LocalEmbeddings()
        vectors = FAISS.from_documents(chunked_docs, embeddings)
        faiss_file_dir = os.path.join(faiss_indices_dir, os.path.splitext(full_filename)[0])
        os.makedirs(faiss_file_dir, exist_ok=True)
        vectors.save_local(faiss_file_dir)
        logging.info(f"FAISS index saved at: {faiss_file_dir}")

        config_file = os.path.join(token_dir, "config.json")
        config = {"word_limit": int(word_limit)}
        with open(config_file, "w") as f:
            json.dump(config, f)
        logging.info(f"Word limit {word_limit} saved at: {config_file}")

        return jsonify({"status": "success", "message": "File processed successfully"})
    except Exception as e:
        logging.error(f"File processing failed: {e}")
        return jsonify({"status": "error", "message": f"Failed to process file: {str(e)}"})


@app.route('/delete_pdf', methods=['POST'])
def delete_pdf():
    data = request.get_json()
    user_code = data.get('User_Code')
    token = data.get('Token')
    pdf_name = data.get('pdf_name')

    if not pdf_name:
        return jsonify({"status": "error", "message": "pdf_name is required"}), 400

    base_dir = "./clients"
    deleted_files = 0
    deleted_indexes = 0

    try:
        user_search_paths = [os.path.join(base_dir, user_code)] if user_code else [os.path.join(base_dir, d) for d in os.listdir(base_dir) if os.path.isdir(os.path.join(base_dir, d))]
        for user_path in user_search_paths:
            if not os.path.isdir(user_path):
                continue
            token_search_paths = [os.path.join(user_path, token)] if token else [os.path.join(user_path, d) for d in os.listdir(user_path) if os.path.isdir(os.path.join(user_path, d))]
            for token_path in token_search_paths:
                if not os.path.isdir(token_path):
                    continue
                file_dir = os.path.join(token_path, "pdfs")
                faiss_dir = os.path.join(token_path, "faiss_indices")
                target_file = os.path.join(file_dir, pdf_name)
                if os.path.isfile(target_file):
                    os.remove(target_file)
                    deleted_files += 1
                    faiss_file_dir = os.path.join(faiss_dir, os.path.splitext(pdf_name)[0])
                    if os.path.exists(faiss_file_dir):
                        for f in os.listdir(faiss_file_dir):
                            os.remove(os.path.join(faiss_file_dir, f))
                        os.rmdir(faiss_file_dir)
                        deleted_indexes += 1
        return jsonify({"status": "success", "message": "Deletion completed", "deleted_files": deleted_files, "deleted_indexes": deleted_indexes})
    except Exception as e:
        logging.error(f"Deletion failed: {e}")
        return jsonify({"status": "error", "message": f"Deletion failed: {str(e)}"})

@app.route('/ask_question', methods=['POST'])
def ask_question():
    data = request.get_json(force=True)
    user_code = data.get("User_Code")
    token = data.get("Token")
    user_token = data.get("User_Token")
    use_external = data.get("Condition", False)
    question = data.get("Question")

    if not all([user_code, token, user_token, question]):
        return jsonify({"status": "error", "message": "Missing required fields", "answer": ""})

    faiss_indices_dir = os.path.join("./clients", user_code, token, "faiss_indices")
    session_key = f"{user_code}_{token}_{user_token}"
    history_dict = strict_conversation_history if use_external else flexible_conversation_history

    if not os.path.exists(faiss_indices_dir) or not os.listdir(faiss_indices_dir):
        logging.info(f"No FAISS indices found for {user_code}, using model fallback.")
        if session_key not in history_dict:
            history_dict[session_key] = []
        history = history_dict[session_key][-10:]
        formatted_history = "\n".join(history)
        if len(history_dict[session_key]) > 10:
            # Summarize using Groq API
            summary_prompt = f"Summarize this conversation:\n{' '.join(history_dict[session_key][:-10])}"
            try:
                completion = client.chat.completions.create(
                    model="meta-llama/llama-4-scout-17b-16e-instruct",
                    messages=[{"role": "user", "content": summary_prompt}],
                    temperature=1,
                    max_completion_tokens=1024,
                    top_p=1,
                    stream=True,
                    stop=None,
                )
                summary = ""
                for chunk in completion:
                    if chunk.choices[0].delta.content:
                        summary += chunk.choices[0].delta.content
                history_dict[session_key] = [f"Summary: {summary}"] + history_dict[session_key][-10:]
                formatted_history = f"Summary: {summary}\n" + "\n".join(history_dict[session_key][-10:])
            except Exception as e:
                logging.error(f"Summary generation failed: {e}")
                return jsonify({"status": "error", "message": f"Summary generation failed: {str(e)}", "answer": ""})

        prompt_template = get_prompt_template(is_strict=False)
        prompt = prompt_template.format(context="", history=formatted_history, input=question)
        try:
            # Invoke Groq API for fallback answer
            completion = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[{"role": "user", "content": prompt}],
                temperature=1,
                max_completion_tokens=1024,
                top_p=1,
                stream=True,
                stop=None,
            )
            answer = ""
            for chunk in completion:
                if chunk.choices[0].delta.content:
                    answer += chunk.choices[0].delta.content
            answer = re.sub(r'^According to[^,.]+[,.]\s*', '', answer).strip()
            history_dict[session_key].extend([f"User: {question}", f"Assistant: {answer}"])
            return jsonify({"status": "success", "message": "Answered using model", "answer": answer})
        except Exception as e:
            logging.error(f"Model invocation failed: {e}")
            return jsonify({"status": "error", "message": f"Failed to process: {str(e)}", "answer": ""})

    try:
        if session_key not in history_dict:
            history_dict[session_key] = []
        history = history_dict[session_key][-2:]
        formatted_history = "\n".join(history)

        embeddings = LocalEmbeddings()
        question_embedding = embeddings.embed_query(question)

        index_scores = []
        for index_dir in os.listdir(faiss_indices_dir):
            index_path = os.path.join(faiss_indices_dir, index_dir)
            if os.path.isdir(index_path):
                try:
                    vectors = FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
                    docs_with_scores = vectors.similarity_search_with_score(question, k=1)
                    if docs_with_scores:
                        score = docs_with_scores[0][1]
                        index_scores.append({'index': vectors, 'score': score, 'name': index_dir})
                except Exception as e:
                    logging.error(f"Error loading {index_dir}: {e}")
                    continue

        if not index_scores:
            return jsonify({"status": "error", "message": "No valid indices", "answer": ""})

        sorted_indices = sorted(index_scores, key=lambda x: x['score'])
        max_indices_to_use = 3
        score_threshold = sorted_indices[0]['score'] * 1.1
        relevant_indices = [idx for idx in sorted_indices if idx['score'] <= score_threshold][:max_indices_to_use]

        if not relevant_indices:
            return jsonify({"status": "error", "message": "No relevant indices", "answer": ""})

        combined_context = []
        sources_used = []
        for idx in relevant_indices:
            vectors = idx['index']
            docs_with_scores = vectors.similarity_search_with_score(question, k=5)
            doc_texts = [doc[0].page_content for doc in docs_with_scores]
            tfidf_vectorizer = TfidfVectorizer(stop_words='english')
            tfidf_matrix = tfidf_vectorizer.fit_transform([question] + doc_texts)
            tfidf_scores = tfidf_matrix[0].dot(tfidf_matrix[1:].T).toarray().flatten()

            weight_vector = 0.5 if idx['name'].endswith(('.csv', '.xlsx')) else 0.7
            weight_tfidf = 0.5 if idx['name'].endswith(('.csv', '.xlsx')) else 0.3
            hybrid_scores = [(doc, weight_vector * (1 - vector_score) + weight_tfidf * tfidf_scores[i]) 
                            for i, (doc, vector_score) in enumerate(docs_with_scores)]
            query_doc_pairs = [(question, doc.page_content) for doc, _ in hybrid_scores]
            
            if query_doc_pairs:
                rerank_scores = cross_encoder.predict(query_doc_pairs)
                reranked_docs = sorted(zip(hybrid_scores, rerank_scores), key=lambda x: x[1], reverse=True)
                top_docs = [doc[0][0] for doc in reranked_docs[:5]]
            else:
                top_docs = [doc[0] for doc in docs_with_scores[:5]]

            combined_context.extend([doc.page_content for doc in top_docs])
            sources_used.append(idx['name'])

        context = "\n".join(combined_context)
        prompt_template = get_prompt_template(use_external)
        prompt = prompt_template.format(context=context, history=formatted_history, input=question)

        # Invoke Groq API with context
        try:
            completion = client.chat.completions.create(
                model="meta-llama/llama-4-scout-17b-16e-instruct",
                messages=[{"role": "user", "content": prompt}],
                temperature=1,
                max_completion_tokens=1024,
                top_p=1,
                stream=True,
                stop=None,
            )
            answer = ""
            for chunk in completion:
                if chunk.choices[0].delta.content:
                    answer += chunk.choices[0].delta.content
            answer = re.sub(r'^According to[^,.]+[,.]\s*', '', answer).strip()
        except Exception as e:
            logging.error(f"Groq API invocation failed: {e}")
            return jsonify({"status": "error", "message": f"Failed to process: {str(e)}", "answer": ""})

        history_dict[session_key].extend([f"User: {question}", f"Assistant: {answer}"])
        if len(history_dict[session_key]) > 10:
            history_dict[session_key] = history_dict[session_key][-10:]

        return jsonify({"status": "success", "message": "", "answer": answer, "sources": sources_used})
    except Exception as e:
        logging.error(f"Error processing question: {e}")
        return jsonify({"status": "error", "message": f"Failed to process: {str(e)}", "answer": ""})
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5011)
